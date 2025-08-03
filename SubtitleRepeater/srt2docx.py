import os
import tkinter as tk
from tkinter import filedialog
from docx import Document


def srt_to_docx(files, suffix):
    for file_path in files:
        document = Document()
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                document.add_paragraph(line.strip())

        base, _ = os.path.splitext(file_path)
        new_path = f"{base}_{suffix}.docx" if suffix else f"{base}.docx"
        document.save(new_path)
    status_label.config(text=f"转换完成：{len(files)} 个 SRT → DOCX")


def docx_to_srt(files, suffix):
    for file_path in files:
        document = Document(file_path)
        base, _ = os.path.splitext(file_path)
        new_path = f"{base}_{suffix}.srt" if suffix else f"{base}.srt"
        with open(new_path, 'w', encoding='utf-8') as f:
            for para in document.paragraphs:
                f.write(para.text + '\n')
    status_label.config(text=f"转换完成：{len(files)} 个 DOCX → SRT")


def add_suffix(files, suffix):
    for file_path in files:
        base, ext = os.path.splitext(file_path)
        new_name = f"{base}_{suffix}{ext}" if suffix else file_path
        os.rename(file_path, new_name)
    status_label.config(text=f"已添加后缀：_{suffix} 到 {len(files)} 个文件")


def remove_suffix(files, suffix):
    for file_path in files:
        base, ext = os.path.splitext(file_path)
        suffix_pattern = f"_{suffix}"
        if base.endswith(suffix_pattern):
            new_name = base[: -len(suffix_pattern)] + ext
            os.rename(file_path, new_name)
    status_label.config(text=f"已移除后缀：_{suffix} 从 {len(files)} 个文件")


def select_files_and_convert(conversion_func):
    files = filedialog.askopenfilenames(filetypes=[("All files", "*.*")])
    suffix = suffix_entry.get().strip()
    if files:
        conversion_func(files, suffix)


# GUI
root = tk.Tk()
root.title("SubSwitch - 字幕格式/文件名转换工具")

tk.Label(root, text="后缀名（可选，例如 ch）").pack(pady=5)
suffix_entry = tk.Entry(root, width=30)
suffix_entry.pack(pady=5)

btn_frame = tk.Frame(root)
btn_frame.pack(pady=10)

tk.Button(btn_frame, text="SRT → DOCX", width=20, command=lambda: select_files_and_convert(srt_to_docx)).grid(
    row=0, column=0, padx=5, pady=5
)
tk.Button(btn_frame, text="DOCX → SRT", width=20, command=lambda: select_files_and_convert(docx_to_srt)).grid(
    row=0, column=1, padx=5, pady=5
)
tk.Button(btn_frame, text="添加后缀名", width=20, command=lambda: select_files_and_convert(add_suffix)).grid(
    row=1, column=0, padx=5, pady=5
)
tk.Button(btn_frame, text="移除后缀名", width=20, command=lambda: select_files_and_convert(remove_suffix)).grid(
    row=1, column=1, padx=5, pady=5
)

status_label = tk.Label(root, text="")
status_label.pack(pady=10)

root.mainloop()
